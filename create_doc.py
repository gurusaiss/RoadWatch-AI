"""
RoadWatch DOC Generator
Creates the Word document submission: packages used + technical assumptions.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Colour helpers ───────────────────────────────────────────────────────

def hex_to_rgb(h):
    h = h.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    tcPr.append(shd)

def set_para_border_bottom(para, color='0D1B2A', sz=12):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

# ── Document setup ────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── Title block ──────────────────────────────────────────────────────────

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run('RoadWatch')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x0D, 0x6E, 0xFF)

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub_para.add_run('AI-Powered Road Safety & Transparency Platform')
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

track_para = doc.add_paragraph()
track_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = track_para.add_run('IIT Madras National Road Safety Hackathon 2026  |  RoadWatch Track')
run3.font.size = Pt(11)
run3.italic = True
run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

set_para_border_bottom(track_para, '1A6BB5', 16)
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 1: Software Packages
# ════════════════════════════════════════════════════════════════════════════

h1 = doc.add_paragraph()
r = h1.add_run('1.  Software Packages & Dependencies')
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x0D, 0x6E, 0xFF)
set_para_border_bottom(h1, '0D6EFF', 10)
doc.add_paragraph()

# ── 1a. Backend ──────────────────────────────────────────────────────────

h2 = doc.add_paragraph()
r = h2.add_run('1.1  Backend (Python 3.10+)')
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x15, 0x52, 0xA0)

backend_packages = [
    ("fastapi", "0.104.1",
     "Core web framework. Provides async REST API with automatic OpenAPI docs."),
    ("uvicorn[standard]", "0.24.0",
     "ASGI server for running FastAPI. 'standard' extras include websocket support."),
    ("sqlalchemy", "2.0.23",
     "ORM for database access. Defines Road, Complaint, Authority, BudgetRecord models."),
    ("pydantic", "2.5.2",
     "Request/response validation and serialisation (v2 with model_validator support)."),
    ("python-multipart", "0.0.6",
     "Required by FastAPI to handle multipart/form-data for file uploads."),
    ("aiofiles", "23.2.1",
     "Async file I/O for saving uploaded complaint images without blocking the event loop."),
    ("httpx", "0.25.2",
     "Async HTTP client used internally by FastAPI's TestClient and AI engine calls."),
    ("pillow", "10.1.0",
     "Image processing: validates uploaded images and converts formats before AI assessment."),
    ("python-dotenv", "1.0.0",
     "Loads .env file for GROQ_API_KEY and OPENAI_API_KEY at runtime."),
    ("groq", "0.4.2",
     "Official Groq SDK for LLaMA 3.3-70b-versatile model. Primary AI engine for chatbot."),
    ("openai", "1.6.1",
     "OpenAI SDK: GPT-4 fallback for chatbot; GPT-4 Vision for damage image assessment."),
    ("geopy", "2.4.1",
     "Geocoding utilities (standby). Haversine distance computation for nearby-roads search."),
    ("haversine", "2.8.0",
     "Fast great-circle distance calculation between GPS coordinates (lat/lon pairs)."),
    ("sqlite3", "stdlib",
     "Bundled with Python. Database engine for roadwatch.db (production: swap to PostgreSQL)."),
]

doc.add_paragraph()
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

# header
hdr = tbl.rows[0].cells
for cell, text, w in zip(hdr,
    ['Package', 'Version', 'Purpose'],
    [Inches(1.8), Inches(0.9), Inches(4.5)]):
    set_cell_bg(cell, '1A4B8A')
    cell.width = w
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10.5)

for i, (pkg, ver, purpose) in enumerate(backend_packages):
    row = tbl.add_row().cells
    bg = 'F0F4F8' if i % 2 == 0 else 'FFFFFF'
    for cell in row:
        set_cell_bg(cell, bg)
    for cell, text, align in zip(row,
        [pkg, ver, purpose],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]):
        p = cell.paragraphs[0]
        p.alignment = align
        r = p.add_run(text)
        r.font.size = Pt(10)
        if cell == row[0]:
            r.bold = True
            r.font.color.rgb = RGBColor(0x0D, 0x4A, 0x99)

doc.add_paragraph()

# ── 1b. Frontend ─────────────────────────────────────────────────────────

h2b = doc.add_paragraph()
r = h2b.add_run('1.2  Frontend (Browser / Vanilla JS)')
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x15, 0x52, 0xA0)

frontend_packages = [
    ("Leaflet.js", "1.9.4",
     "Interactive map rendering. Road markers, popups, tile layers (OpenStreetMap)."),
    ("Leaflet.MarkerCluster", "1.5.3",
     "Clusters nearby map markers to prevent overlapping at low zoom levels."),
    ("Chart.js", "4.4.0",
     "Renders doughnut, bar, and pie charts for analytics dashboard."),
    ("marked.js", "9.1.6",
     "Parses Markdown to HTML for rich chatbot message rendering."),
    ("Service Worker API", "Browser native",
     "PWA offline caching: caches /api/roads, /api/stats, and all static assets."),
    ("Fetch API", "Browser native",
     "All HTTP calls to the FastAPI backend (JSON + multipart)."),
    ("Geolocation API", "Browser native",
     "GPS coordinates for 'Near Me' map feature and complaint GPS auto-fill."),
    ("OpenStreetMap tiles", "Free CDN",
     "Base map tiles served from tile.openstreetmap.org. No API key required."),
]

doc.add_paragraph()
tbl2 = doc.add_table(rows=1, cols=3)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr2 = tbl2.rows[0].cells
for cell, text in zip(hdr2, ['Library / API', 'Version', 'Purpose']):
    set_cell_bg(cell, '1A4B8A')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10.5)

for i, (pkg, ver, purpose) in enumerate(frontend_packages):
    row = tbl2.add_row().cells
    bg = 'F0F4F8' if i % 2 == 0 else 'FFFFFF'
    for cell in row:
        set_cell_bg(cell, bg)
    for cell, text, align in zip(row,
        [pkg, ver, purpose],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]):
        p = cell.paragraphs[0]
        p.alignment = align
        r = p.add_run(text)
        r.font.size = Pt(10)
        if cell == row[0]:
            r.bold = True
            r.font.color.rgb = RGBColor(0x0D, 0x4A, 0x99)

doc.add_paragraph()

# ── 1c. External AI Services ─────────────────────────────────────────────

h2c = doc.add_paragraph()
r = h2c.add_run('1.3  External AI / Cloud Services (Optional Keys)')
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x15, 0x52, 0xA0)

ai_services = [
    ("Groq Cloud API", "LLaMA 3.3-70b-versatile",
     "Primary LLM engine for chatbot. Free tier available. Set GROQ_API_KEY in .env."),
    ("OpenAI API", "GPT-4 / GPT-4 Vision",
     "Fallback LLM + damage image assessment. Set OPENAI_API_KEY in .env."),
    ("Rule-based fallback", "Built-in",
     "Fully offline fallback if no API keys provided. No external calls made."),
]

doc.add_paragraph()
tbl3 = doc.add_table(rows=1, cols=3)
tbl3.style = 'Table Grid'
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr3 = tbl3.rows[0].cells
for cell, text in zip(hdr3, ['Service', 'Model Used', 'Notes']):
    set_cell_bg(cell, '1A4B8A')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10.5)

for i, (svc, model, note) in enumerate(ai_services):
    row = tbl3.add_row().cells
    bg = 'F0F4F8' if i % 2 == 0 else 'FFFFFF'
    for cell in row:
        set_cell_bg(cell, bg)
    for cell, text, align in zip(row,
        [svc, model, note],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]):
        p = cell.paragraphs[0]
        p.alignment = align
        r = p.add_run(text)
        r.font.size = Pt(10)
        if cell == row[0]:
            r.bold = True
            r.font.color.rgb = RGBColor(0x0D, 0x4A, 0x99)

doc.add_paragraph()
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 2: Technical Assumptions
# ════════════════════════════════════════════════════════════════════════════

h1b = doc.add_paragraph()
r = h1b.add_run('2.  Technical Assumptions')
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x0D, 0x6E, 0xFF)
set_para_border_bottom(h1b, '0D6EFF', 10)
doc.add_paragraph()

assumptions = [
    # (category, heading, body)
    ("Data & Database",
     "Seed data represents realistic but synthesised road records",
     "All 35 road records (NH-44, NH-48, NH-16, M1-UK, I-95-US, etc.) use publicly verifiable "
     "road names, approximate lengths, and realistic construction dates across 13 countries. "
     "Contractor names, Executive Engineer names, and exact budget figures are "
     "illustrative/synthesised for demonstration purposes. In a production system these would "
     "be sourced directly from MoRTH/NRIDA/PMGSY APIs or official data exports."),

    ("Data & Database",
     "SQLite used for development; PostgreSQL recommended for production",
     "SQLite with check_same_thread=False is used for simplicity during the hackathon. "
     "The SQLAlchemy ORM layer is database-agnostic — switching to PostgreSQL requires only "
     "a one-line DATABASE_URL change in database.py with no model changes needed."),

    ("Data & Database",
     "Road IDs follow official Indian and international naming conventions",
     "Indian roads use codes like NH-44, SH-36-MH, MDR-23-TN following NHAI / MoRTH "
     "classification. International roads use country-specific conventions (M1-UK for UK "
     "motorways, I-95-US for US Interstate, AH1-BD for Asian Highway, etc.)."),

    ("AI & Machine Learning",
     "Damage assessment is rule-based mock unless OpenAI Vision key is provided",
     "Without OPENAI_API_KEY, the /api/complaints endpoint returns a structured mock "
     "assessment with randomised damage type (pothole/crack/subsidence/erosion/waterlogging) "
     "and severity. Confidence score is set to 0.85 for mock responses. With a valid key, "
     "GPT-4 Vision analyses the actual uploaded image."),

    ("AI & Machine Learning",
     "LLM chatbot falls back gracefully without API keys",
     "The AI engine tries: (1) Groq LLaMA-3.3-70b-versatile, (2) OpenAI GPT-4, "
     "(3) comprehensive rule-based engine with 40+ query handlers. The rule-based fallback "
     "handles all core queries (road conditions, budget info, critical roads list, anomalies, "
     "contractor details, complaint routing) using live data fetched from the database, "
     "ensuring full functionality even offline."),

    ("Geolocation & Mapping",
     "GPS coordinates in seed data are approximate road centroids",
     "lat_center / lon_center coordinates are approximate midpoints of the named roads "
     "using publicly available highway GPS data. They are accurate enough for map display "
     "and nearby-road calculations but not survey-grade."),

    ("Geolocation & Mapping",
     "Nearby roads search uses flat-earth Haversine approximation",
     "The /api/roads/nearby endpoint uses the haversine formula assuming a spherical Earth "
     "(R = 6371 km). This introduces < 0.5% error at distances under 100 km, which is "
     "acceptable for road proximity search."),

    ("Complaint Routing",
     "Authority routing priority: database record > rule-based fallback",
     "When a complaint is filed, the system first queries the Authority table for a matching "
     "(country, state, road_type, district) record. If no exact match is found, it falls "
     "back to a hard-coded ROUTING_RULES dictionary covering all 7 countries. "
     "Production deployment would populate the Authority table from official department directories."),

    ("Complaint Routing",
     "Budget anomaly threshold set at 65% utilisation",
     "A road's budget is flagged as anomalous if amount_spent / amount_sanctioned < 0.65. "
     "This threshold is based on general public works norms suggesting that under-utilisation "
     "below 65% warrants scrutiny. The threshold is configurable in main.py."),

    ("Offline / PWA",
     "Service Worker caches last-fetched road and stats data",
     "The PWA service worker uses a network-first strategy for API routes: it always tries "
     "the network first and falls back to the last cached response if offline. Static assets "
     "(HTML, CSS, JS) use cache-first. Cache version is 'roadwatch-v1'; bumping this string "
     "forces a full refresh on next visit."),

    ("Security & Privacy",
     "No authentication implemented (hackathon scope)",
     "The current implementation has no user authentication or role-based access control. "
     "For production: JWT-based auth for complaint status updates, OAuth for authority logins, "
     "and row-level security on the complaint table would be added."),

    ("Internationalisation",
     "Currency display is informational; no conversion performed",
     "Budget figures are stored with their original currency (INR, GBP, USD, BDT, ZAR, MGA, PLN). "
     "The UI displays the stored currency symbol without live FX conversion. "
     "All amounts are in local currency as reported by the respective road authority."),

    ("Image Uploads",
     "Uploaded images are stored locally in data/uploads/",
     "In the hackathon setup, images are saved to ../data/uploads/ relative to the backend. "
     "File size is limited to 10 MB per upload. Accepted formats: JPEG, PNG, WebP, GIF. "
     "Production would use cloud object storage (AWS S3 / Google Cloud Storage)."),

    ("Road Classification",
     "India road type hierarchy follows MoRTH official classification",
     "NH (National Highway) -> NHAI/MoRTH, SH (State Highway) -> State PWD, "
     "MDR (Major District Road) -> District PWD, ODR (Other District Road) -> District PWD, "
     "VR (Village Road) -> Municipal/Gram Panchayat. "
     "Expressways are treated as a special sub-type of NH for routing purposes."),
]

# Group by category
from collections import OrderedDict
by_cat = OrderedDict()
for cat, heading, body in assumptions:
    by_cat.setdefault(cat, []).append((heading, body))

ass_num = 1
for cat, items in by_cat.items():
    # Category heading
    cat_para = doc.add_paragraph()
    r = cat_para.add_run(f'  {cat}')
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x15, 0x52, 0xA0)
    set_cell_bg  # not applicable here
    doc.add_paragraph()

    for heading, body in items:
        # Assumption heading
        h_para = doc.add_paragraph()
        h_para.paragraph_format.left_indent = Inches(0.3)
        r1 = h_para.add_run(f'A{ass_num}.  ')
        r1.bold = True
        r1.font.size = Pt(11.5)
        r1.font.color.rgb = RGBColor(0xFF, 0x6B, 0x35)
        r2 = h_para.add_run(heading)
        r2.bold = True
        r2.font.size = Pt(11.5)
        r2.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        # Body
        b_para = doc.add_paragraph()
        b_para.paragraph_format.left_indent = Inches(0.6)
        b_para.paragraph_format.space_after = Pt(8)
        r3 = b_para.add_run(body)
        r3.font.size = Pt(10.5)
        r3.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        ass_num += 1

    doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 3: Evaluation Criteria Coverage
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()

h1c = doc.add_paragraph()
r = h1c.add_run('3.  Evaluation Criteria Coverage')
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x0D, 0x6E, 0xFF)
set_para_border_bottom(h1c, '0D6EFF', 10)
doc.add_paragraph()

criteria_rows = [
    ("Data Accuracy",
     "Road Type, last relaying date, contractor name, EE stored per road. "
     "All fields have source_label (NHAI, State PWD, PMGSY, etc.) and source_url. "
     "Budget records include verified flag and source link."),
    ("Correct Complaint Routing",
     "Complaint filing wizard identifies road type -> queries Authority table -> "
     "returns correct Executive Engineer name, email, phone, department. "
     "AI engine (Groq/GPT-4/fallback) confirms routing in chatbot."),
    ("Budget Transparency with Source",
     "Each road has budget_sanctioned, budget_spent, financial_year, currency. "
     "BudgetRecord table stores multi-year history with contractor and source_url. "
     "Anomaly flag auto-set when utilisation < 65%."),
    ("Global / Cross-Country Applicability",
     "13 countries seeded: India, UK, USA, Bangladesh, South Africa, Madagascar, Poland, "
     "Germany, Australia, Nigeria, Kenya, Japan, Malaysia. "
     "37 authorities mapped across all countries. Currency displayed in local denomination."),
    ("Offline Functionality",
     "Service Worker (service-worker.js) caches /api/roads, /api/stats, and all frontend "
     "assets. App works fully offline for road browsing. Complaints queue locally (future)."),
    ("UI & Accessibility",
     "Responsive CSS (Grid/Flexbox), hamburger nav at 768px, colour-blind-safe palette, "
     "ARIA labels on interactive elements, keyboard-navigable wizard, toast notifications."),
]

tbl4 = doc.add_table(rows=1, cols=2)
tbl4.style = 'Table Grid'
tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr4 = tbl4.rows[0].cells
for cell, text in zip(hdr4, ['Evaluation Criterion', 'How RoadWatch Satisfies It']):
    set_cell_bg(cell, '0D6EFF')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(11)

for i, (crit, how) in enumerate(criteria_rows):
    row = tbl4.add_row().cells
    bg = 'EAF3FF' if i % 2 == 0 else 'FFFFFF'
    set_cell_bg(row[0], bg)
    set_cell_bg(row[1], bg)
    p0 = row[0].paragraphs[0]
    r = p0.add_run(f'[checkmark] {crit}')
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(0x0D, 0x6E, 0xFF)

    p1 = row[1].paragraphs[0]
    r2 = p1.add_run(how)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

doc.add_paragraph()

# ── Footer note ──────────────────────────────────────────────────────────
doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_border_bottom(foot, 'CCCCCC', 8)

foot2 = doc.add_paragraph()
foot2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot2.add_run(
    'RoadWatch  |  IIT Madras National Road Safety Hackathon 2026  |  '
    'All packages are open-source or have free-tier API access  |  '
    'No proprietary paid services required to run the demo')
r.font.size = Pt(9)
r.italic = True
r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

# ── Save ─────────────────────────────────────────────────────────────────
out_path = r"D:\C\CODING\Hackathon\Road Safety\roadwatch\RoadWatch_Packages_Assumptions.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
